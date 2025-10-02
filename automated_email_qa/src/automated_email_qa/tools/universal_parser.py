"""
Universal Document Parser for Email QA System
=============================================
Purpose: Parse various document formats to extract email requirements
Author: Rishabh Sharma
Date: 2024

This module handles parsing of copy documents in multiple formats to extract:
- Subject lines and preview text
- From name and email addresses
- CTAs and their destinations
- Content requirements and segments
- Special instructions

Metadata:
---------
Classes:
    UniversalDocumentParser: Main parser class for all document types
    
Key Data Types:
    content (Union[str, bytes]): Raw document content
        - str: Text or HTML content
        - bytes: Binary content requiring decoding
    
    filename (Optional[str]): Original filename for format detection
    
    parsed_requirements (Dict[str, Any]): Extracted requirements
        - subject_lines (List[str]): All subject line variations
        - preview_text (str): Preview/preheader text
        - from_name (str): Sender name
        - from_email (str): Sender email address
        - ctas (List[Dict[str, str]]): CTA buttons/links
            - text (str): CTA button text
            - link (str): CTA destination URL
        - links (List[str]): All URLs found in document
        - segments (Dict[str, Dict]): Segment-specific requirements
        - content_modules (List[str]): Required content sections
        - special_notes (List[str]): Additional instructions
        - encoding_fixed (bool): Whether encoding fixes were applied

Integration with dynamic_rules.py:
    - Extracted 'ctas' matches rules['links'] structure
    - Extracted 'segments' aligns with rules['segments']
    - CTA text casing will be validated against rules['brand']['cta']['style']
"""

import re
import email
from email import policy
from email.parser import Parser
from bs4 import BeautifulSoup
from typing import Dict, Any, List, Optional, Union
import chardet
import logging
from pathlib import Path

import docx  # python-docx
import PyPDF2  # pypdf2
import openpyxl  # openpyxl
import io
from docx import Document
from PyPDF2 import PdfReader

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class UniversalDocumentParser:
    """
    Parses various document formats to extract email requirements.
    
    This class handles multiple file formats and extracts structured
    requirements that can be validated against dynamic rules.
    
    Supported formats:
        - Plain text (.txt)
        - HTML (.html)
        - Email files (.eml)
        - Future: DOCX, PDF, XLSX
    """
    
    # Encoding replacements from document specification
    ENCODING_FIXES: Dict[str, str] = {
        'â€™': "'",      # Apostrophe
        'â€"': "-",      # Em dash
        'â€œ': '"',      # Left double quote
        'â€': '"',       # Right double quote
        'â€¦': '...',    # Ellipsis
        'â€¢': '•',      # Bullet
        'â€˜': "'",      # Left single quote
        'Ã©': 'é',       # e acute
        'Ã¨': 'è',       # e grave
        'Ã ': 'à',       # a grave
        '&nbsp;': ' ',   # Non-breaking space
        '&amp;': '&',    # Ampersand
        '&lt;': '<',     # Less than
        '&gt;': '>',     # Greater than
        '&#39;': "'",    # Apostrophe (HTML entity)
        '&quot;': '"',   # Quote (HTML entity)
    }
    
    def __init__(self):
        """Initialize the Universal Document Parser."""
        self.encoding_issues_found: List[str] = []
        logger.info("Initialized UniversalDocumentParser")
    
    def parse_document(self, 
                    content: Union[str, bytes], 
                    filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse any document format to extract email requirements.
        
        Args:
            content (Union[str, bytes]): Document content to parse
            filename (Optional[str]): Original filename for format detection
            
        Returns:
            Dict[str, Any]: Extracted requirements structure
        """
        # Reset tracking
        self.encoding_issues_found = []
        
        # CRITICAL: Handle binary formats BEFORE any text decoding
        if isinstance(content, bytes):
            if filename:
                file_ext = Path(filename).suffix.lower()
                
                # Handle binary document formats directly
                if file_ext == '.docx':
                    return self._parse_docx_binary(content)
                elif file_ext == '.pdf':
                    return self._parse_pdf_binary(content)
                elif file_ext in ['.xlsx', '.xls']:
                    return self._parse_excel_binary(content)
            
            # Only for text formats - detect and decode
            try:
                detected = chardet.detect(content)
                encoding = detected.get('encoding', 'utf-8')
                confidence = detected.get('confidence', 0)
                
                logger.info(f"Detected encoding: {encoding} (confidence: {confidence:.2%})")
                
                content = content.decode(encoding, errors='replace')
            except (UnicodeDecodeError, TypeError):
                content = content.decode('utf-8', errors='replace')
                logger.warning("Failed to decode, using UTF-8 fallback")
        
        # Size check after decoding
        if len(content) > 500000:  # 500KB text limit
            logger.warning(f"Large document ({len(content)} chars), truncating")
            content = content[:500000] + "\n[TRUNCATED - Document exceeded size limit]"
        
        # Fix encoding issues in text
        content = self._fix_encoding(content)
        
        # Text-based format detection and parsing
        if filename:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.eml':
                return self._parse_eml_format(content)
            elif file_ext == '.html':
                return self._parse_html_requirements(content)
        
        # Auto-detect format from content
        if self._is_html(content):
            return self._parse_html_requirements(content)
        elif self._is_eml(content):
            return self._parse_eml_format(content)
        else:
            return self._parse_text_requirements(content)

    def _parse_docx_binary(self, content: bytes) -> Dict[str, Any]:
        """Extract requirements from DOCX binary content."""
        try:
            # Create document from bytes
            doc = Document(io.BytesIO(content))
            
            # Extract all text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            # Join all text and parse as requirements
            full_text = '\n'.join(text_content)
            
            if len(full_text) > 500000:  # Size limit
                logger.warning(f"Large DOCX content ({len(full_text)} chars), truncating")
                full_text = full_text[:500000] + "\n[TRUNCATED]"
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return self._parse_text_requirements(full_text)
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return {
                'subject_lines': [],
                'preview_text': '',
                'from_name': '',
                'from_email': '',
                'ctas': [],
                'links': [],
                'segments': {},
                'content_modules': [],
                'special_notes': [f"DOCX parsing failed: {str(e)}"],
                'encoding_fixed': False,
                'encoding_issues': []
            }

    def _parse_pdf_binary(self, content: bytes) -> Dict[str, Any]:
        """Extract requirements from PDF binary content."""
        try:
            pdf_reader = PdfReader(io.BytesIO(content))
            
            text_content = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text.strip())
            
            full_text = '\n'.join(text_content)
            
            if len(full_text) > 500000:  # Size limit
                logger.warning(f"Large PDF content ({len(full_text)} chars), truncating")
                full_text = full_text[:500000] + "\n[TRUNCATED]"
            
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return self._parse_text_requirements(full_text)
            
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            return {
                'subject_lines': [],
                'preview_text': '',
                'from_name': '',
                'from_email': '',
                'ctas': [],
                'links': [],
                'segments': {},
                'content_modules': [],
                'special_notes': [f"PDF parsing failed: {str(e)}"],
                'encoding_fixed': False,
                'encoding_issues': []
            }

    def _parse_excel_binary(self, content: bytes) -> Dict[str, Any]:
        """Extract requirements from Excel binary content."""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            
            text_content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Add sheet name as header
                text_content.append(f"\n=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell).strip())
                    
                    if any(row_text):  # Skip empty rows
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_content)
            
            if len(full_text) > 500000:  # Size limit  
                logger.warning(f"Large Excel content ({len(full_text)} chars), truncating")
                full_text = full_text[:500000] + "\n[TRUNCATED]"
            
            logger.info(f"Extracted {len(full_text)} characters from Excel")
            return self._parse_text_requirements(full_text)
            
        except Exception as e:
            logger.error(f"Excel parsing failed: {e}")
            return {
                'subject_lines': [],
                'preview_text': '',
                'from_name': '',
                'from_email': '',
                'ctas': [],
                'links': [],
                'segments': {},
                'content_modules': [],
                'special_notes': [f"Excel parsing failed: {str(e)}"],
                'encoding_fixed': False,
                'encoding_issues': []
            }
    
    def _fix_encoding(self, text: str) -> str:
        """
        Fix common encoding issues in text.
        
        Implements the encoding fixes specified in the document.
        
        Args:
            text (str): Text with potential encoding issues
            
        Returns:
            str: Text with encoding issues fixed
        """
        original_text = text
        
        for broken, fixed in self.ENCODING_FIXES.items():
            if broken in text:
                self.encoding_issues_found.append(f"Fixed: {broken} -> {fixed}")
                text = text.replace(broken, fixed)
        
        if text != original_text:
            logger.info(f"Fixed {len(self.encoding_issues_found)} encoding issues")
        
        return text
    
    def _is_html(self, content: str) -> bool:
        """
        Check if content is HTML format.
        
        Args:
            content (str): Content to check
            
        Returns:
            bool: True if content appears to be HTML
        """
        html_indicators = ['<html', '<body', '<div', '<table', '<!DOCTYPE']
        content_lower = content[:1000].lower()  # Check first 1000 chars
        
        return any(indicator in content_lower for indicator in html_indicators)
    
    def _is_eml(self, content: str) -> bool:
        """
        Check if content is EML email format.
        
        Args:
            content (str): Content to check
            
        Returns:
            bool: True if content appears to be EML
        """
        eml_headers = ['Subject:', 'From:', 'To:', 'Date:', 'Content-Type:', 'MIME-Version:']
        
        # Check for multiple email headers in first 500 chars
        header_count = sum(1 for header in eml_headers if header in content[:500])
        
        return header_count >= 2
    
    def _parse_text_requirements(self, content: str) -> Dict[str, Any]:
        """
        Extract requirements from plain text documents.
        
        Implements pattern matching to find subject lines, CTAs, URLs, etc.
        
        Args:
            content (str): Text content to parse
            
        Returns:
            Dict[str, Any]: Extracted requirements
        """
        requirements: Dict[str, Any] = {
            'subject_lines': [],
            'preview_text': '',
            'from_name': '',
            'from_email': '',
            'ctas': [],
            'links': [],
            'segments': {},
            'content_modules': [],
            'special_notes': [],
            'encoding_fixed': len(self.encoding_issues_found) > 0,
            'encoding_issues': self.encoding_issues_found
        }
        
        lines = content.split('\n')
        current_segment = None
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Extract subject lines (multiple patterns)
            if any(pattern in line_lower for pattern in ['subject:', 'subject line:', 'sl:']):
                if ':' in line:
                    subject = line.split(':', 1)[1].strip()
                    if subject and subject not in requirements['subject_lines']:
                        requirements['subject_lines'].append(subject)
                        logger.debug(f"Found subject line: {subject}")
            
            # Extract preview text
            elif any(pattern in line_lower for pattern in ['preview:', 'preheader:', 'preview text:']):
                if ':' in line:
                    requirements['preview_text'] = line.split(':', 1)[1].strip()
            
            # Extract From details
            elif 'from name:' in line_lower and ':' in line:
                requirements['from_name'] = line.split(':', 1)[1].strip()
            elif 'from email:' in line_lower and ':' in line:
                requirements['from_email'] = line.split(':', 1)[1].strip()
            elif 'from:' in line_lower and '@' in line:
                # Handle combined "From: Name <email@domain.com>"
                from_match = re.search(r'from:\s*(.+?)\s*<(.+?)>', line, re.IGNORECASE)
                if from_match:
                    requirements['from_name'] = from_match.group(1).strip()
                    requirements['from_email'] = from_match.group(2).strip()
            
            # Extract segments
            elif 'segment:' in line_lower and ':' in line:
                current_segment = line.split(':', 1)[1].strip()
                requirements['segments'][current_segment] = {
                    'requirements': [],
                    'line_number': i + 1
                }
            
            # Extract CTAs (uppercase patterns as per document)
            elif re.match(r'^[A-Z][A-Z\s]+[A-Z]$', line_stripped):
                potential_cta = line_stripped
                # CTAs are typically 3-30 characters and 1-5 words
                if 3 <= len(potential_cta) <= 30 and len(potential_cta.split()) <= 5:
                    # Look ahead for URL on next line
                    next_url = ''
                    if i + 1 < len(lines):
                        url_match = re.search(r'https?://[^\s]+', lines[i + 1])
                        if url_match:
                            next_url = url_match.group(0)
                    
                    requirements['ctas'].append({
                        'text': potential_cta,
                        'link': next_url,
                        'line_number': i + 1
                    })
                    logger.debug(f"Found CTA: {potential_cta}")
            
            # Extract URLs
            urls = re.findall(r'https?://[^\s<>"]+', line)
            for url in urls:
                clean_url = url.rstrip('.,;:)')  # Remove trailing punctuation
                if clean_url not in requirements['links']:
                    requirements['links'].append(clean_url)
            
            # Extract content modules
            if 'module:' in line_lower or 'section:' in line_lower:
                if ':' in line:
                    module = line.split(':', 1)[1].strip()
                    if module not in requirements['content_modules']:
                        requirements['content_modules'].append(module)
            
            # Add to current segment if active
            if current_segment and line_stripped and ':' not in line_stripped:
                requirements['segments'][current_segment]['requirements'].append(line_stripped)
            
            # Special notes (lines starting with NOTE:, IMPORTANT:, etc.)
            if any(line_upper.startswith(marker) for marker in ['NOTE:', 'IMPORTANT:', 'ATTENTION:'] 
                   for line_upper in [line_stripped.upper()]):
                requirements['special_notes'].append(line_stripped)
        
        return requirements
    
    def _parse_html_requirements(self, content: str) -> Dict[str, Any]:
        """
        Extract requirements from HTML documents.
        
        Uses BeautifulSoup to parse HTML structure and extract requirements.
        
        Args:
            content (str): HTML content to parse
            
        Returns:
            Dict[str, Any]: Extracted requirements
        """
        soup = BeautifulSoup(content, 'lxml')
        
        # First extract text and parse as text
        text_content = soup.get_text(separator='\n', strip=True)
        requirements = self._parse_text_requirements(text_content)
        
        # Additionally extract HTML-specific elements
        
        # Extract links with context
        for link in soup.find_all('a', href=True):
            link_text = link.get_text(strip=True)
            link_url = link['href']
            
            if link_text and link_url:
                # Check if it's a CTA based on parent element
                parent_classes = ' '.join(link.parent.get('class', []))
                is_cta = any(indicator in parent_classes.lower() 
                           for indicator in ['button', 'cta', 'btn'])
                
                if is_cta or link_text.isupper():
                    # Add as CTA if not already present
                    cta_exists = any(cta['text'] == link_text for cta in requirements['ctas'])
                    if not cta_exists:
                        requirements['ctas'].append({
                            'text': link_text,
                            'link': link_url,
                            'html_class': parent_classes
                        })
        
        # Extract meta information if present
        title_tag = soup.find('title')
        if title_tag and title_tag.text.strip():
            requirements['special_notes'].append(f"Page title: {title_tag.text.strip()}")
        
        return requirements
    
    def _parse_eml_format(self, content: str) -> Dict[str, Any]:
        """
        Extract requirements from EML email format files.
        
        Args:
            content (str): EML content to parse
            
        Returns:
            Dict[str, Any]: Extracted requirements including email components
        """
        msg = email.message_from_string(content, policy=policy.default)
        
        # Start with base requirements
        requirements = self._parse_text_requirements('')
        
        # Extract email headers
        requirements['subject_lines'] = [msg.get('Subject', '')]
        
        from_header = msg.get('From', '')
        if from_header:
            from_match = re.match(r'(.+?)\s*<(.+?)>', from_header)
            if from_match:
                requirements['from_name'] = from_match.group(1).strip()
                requirements['from_email'] = from_match.group(2).strip()
            else:
                requirements['from_email'] = from_header.strip()
        
        # Extract body content
        for part in msg.walk():
            content_type = part.get_content_type()
            
            if content_type == 'text/html':
                html = part.get_payload(decode=True)
                if html:
                    html_text = html.decode('utf-8', errors='ignore')
                    html_requirements = self._parse_html_requirements(html_text)
                    # Merge HTML requirements
                    requirements['ctas'].extend(html_requirements['ctas'])
                    requirements['links'].extend(html_requirements['links'])
            
            elif content_type == 'text/plain':
                text = part.get_payload(decode=True)
                if text:
                    text_content = text.decode('utf-8', errors='ignore')
                    text_requirements = self._parse_text_requirements(text_content)
                    # Merge text requirements
                    for key in ['content_modules', 'special_notes']:
                        requirements[key].extend(text_requirements[key])
        
        # Remove duplicates
        requirements['links'] = list(set(requirements['links']))
        
        return requirements
    
    def extract_email_components(self, email_content: str) -> Dict[str, Any]:
        """
        Extract all components from email content (HTML or EML).
        
        This method specifically handles email content rather than
        requirement documents.
        
        Args:
            email_content (str): Email content to parse
            
        Returns:
            Dict[str, Any]: Extracted email components for validation
        """
        components = {
            'subject': '',
            'from': '',
            'preview_text': '',
            'html_body': '',
            'plain_body': '',
            'headers': {}
        }
        
        if self._is_eml(email_content):
            msg = email.message_from_string(email_content, policy=policy.default)
            
            components['subject'] = msg.get('Subject', '')
            components['from'] = msg.get('From', '')
            
            for part in msg.walk():
                if part.get_content_type() == 'text/html':
                    html = part.get_payload(decode=True)
                    if html:
                        components['html_body'] = html.decode('utf-8', errors='ignore')
                elif part.get_content_type() == 'text/plain':
                    text = part.get_payload(decode=True)
                    if text:
                        components['plain_body'] = text.decode('utf-8', errors='ignore')
        else:
            # Assume HTML
            components['html_body'] = email_content
        
        # Fix encoding in all components
        for key in components:
            if isinstance(components[key], str):
                components[key] = self._fix_encoding(components[key])
        
        return components